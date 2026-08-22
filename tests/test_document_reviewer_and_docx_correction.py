"""
tests/test_document_reviewer_and_docx_correction.py — Document Control final
correction pass: (A) AI-Assisted SOP output must be a genuine DOCX, never
Markdown; (B) the Reviewer/Department Head/Quality Head/Plant Head approval
assignment must resolve by name to an internal user id, never a raw
Supabase user id typed by hand — extending the existing Approver Pool
mechanism (qms_document_database.py's ALL_POOL_ROLES/resolve_pool_approvers)
with a fourth, single-decider "reviewer" role for the 'under_review' step,
mirroring exactly how "department_head"/"quality_head"/"plant_head" already
resolve for the 'effective' step.

Deliberately no autouse app_context fixture (see test_document_quality_
release.py's identical note) — every test drives the document exclusively
through the `client` fixture so a mid-test TenantContext monkeypatch is
actually re-read by conftest.py's before_request shim.
"""

import io

from docx import Document as DocxDocument

from pharmagpt import qms_document_database as qdb


def _create_submittable(client, title="Cleaning SOP"):
    doc = client.post("/qms/documents", json={"title": title, "content": "# Cleaning SOP\n\nSome content."}).get_json()
    did = doc["id"]
    client.post(f"/qms/documents/{did}/self-check")
    r = client.post(f"/qms/documents/{did}/versions/upload",
                     data={"file": (io.BytesIO(b"final content"), "final.txt")},
                     content_type="multipart/form-data")
    assert r.status_code == 201, r.get_json()
    return did


# ── A. AI-assisted SOP DOCX output ───────────────────────────────────────────

def test_export_docx_is_a_valid_openable_docx(client):
    doc = client.post("/qms/documents", json={
        "title": "Cleaning of Tablet Compression Machine",
        "content": "# Cleaning of Tablet Compression Machine\n\n## 1. Purpose\nAI-generated purpose text.",
    }).get_json()
    r = client.post(f"/qms/documents/{doc['id']}/export/docx")
    assert r.status_code == 200
    assert r.mimetype == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    docx_doc = DocxDocument(io.BytesIO(r.data))
    assert docx_doc.paragraphs  # openable and non-empty


def test_export_docx_filename_ends_in_docx_and_is_not_markdown(client):
    doc = client.post("/qms/documents", json={"title": "Calibration SOP", "content": "x"}).get_json()
    r = client.post(f"/qms/documents/{doc['id']}/export/docx")
    cd = r.headers.get("Content-Disposition", "")
    assert ".docx" in cd
    assert ".md" not in cd
    # Genuine DOCX packages are zip archives — raw markdown text is not.
    assert r.data[:2] == b"PK"


def test_export_docx_preserves_ai_generated_content(client):
    distinctive_text = "Operating parameters: compression force 8-12 kN."
    doc = client.post("/qms/documents", json={
        "title": "Cleaning SOP", "content": f"# Cleaning SOP\n\n## 1. Purpose\n{distinctive_text}",
    }).get_json()
    r = client.post(f"/qms/documents/{doc['id']}/export/docx")
    docx_doc = DocxDocument(io.BytesIO(r.data))
    full_text = "\n".join(p.text for p in docx_doc.paragraphs)
    assert distinctive_text in full_text


def test_export_docx_preserves_controlled_template_headings(client):
    structure = [
        {"heading": "1. Purpose", "sub_headings": []},
        {"heading": "2. Scope", "sub_headings": []},
    ]
    t = client.post("/qms/documents/templates",
                     json={"doc_type": "SOP", "name": "Cleaning Template", "structure": structure}).get_json()
    doc = client.post("/qms/documents", json={
        "title": "Cleaning SOP", "template_id": t["id"],
        "content": "# Cleaning SOP\n\n## 1. Purpose\nPurpose text.\n\n## 2. Scope\nScope text.",
    }).get_json()
    r = client.post(f"/qms/documents/{doc['id']}/export/docx")
    docx_doc = DocxDocument(io.BytesIO(r.data))
    full_text = "\n".join(p.text for p in docx_doc.paragraphs)
    assert "1. Purpose" in full_text
    assert "2. Scope" in full_text


# ── B. Approval workflow — Reviewer/Department Head/Quality Head/Plant Head ─

def test_all_pool_roles_includes_reviewer():
    assert set(qdb.ALL_POOL_ROLES) == {"reviewer", "department_head", "quality_head", "plant_head"}


def test_approver_pool_route_accepts_reviewer_role(client):
    r = client.post("/qms/documents/approver-pool",
                     json={"pool_role": "reviewer", "user_id": "rev-1", "display_name": "Rita Reviewer"})
    assert r.status_code == 201, r.get_json()
    pool = client.get("/qms/documents/approver-pool").get_json()
    assert {"reviewer": "rev-1"}.items() <= {p["pool_role"]: p["user_id"] for p in pool}.items()


def test_approver_pool_route_still_rejects_unknown_role(client):
    r = client.post("/qms/documents/approver-pool", json={"pool_role": "site_head", "user_id": "x"})
    assert r.status_code == 400


def test_approver_pool_route_requires_user_id_for_reviewer(client):
    """Mandatory-role-cannot-be-skipped, at the point it's actually
    enforceable: the pool config save itself refuses an empty assignment —
    same existing rule every pool_role has always had, now also covering
    'reviewer'."""
    r = client.post("/qms/documents/approver-pool", json={"pool_role": "reviewer", "user_id": ""})
    assert r.status_code == 400


def test_resolve_pool_reviewer_returns_none_when_unconfigured(db_path):
    document = {"company_id": "co-1", "department": ""}
    assert qdb.resolve_pool_reviewer(document) is None


def test_resolve_pool_reviewer_returns_configured_entry(db_path):
    qdb.set_approver_pool_member("co-1", "", "reviewer", "rev-9", "Rita")
    entry = qdb.resolve_pool_reviewer({"company_id": "co-1", "department": ""})
    assert entry == {"user_id": "rev-9", "display_name": "Rita", "pool_role": "reviewer"}


def test_start_workflow_auto_assigns_configured_reviewer(client):
    client.post("/qms/documents/approver-pool",
                json={"pool_role": "reviewer", "user_id": "rev-42", "display_name": "Rita Reviewer"})
    did = _create_submittable(client)
    r = client.post(f"/qms/documents/{did}/workflow/start")
    assert r.status_code == 201, r.get_json()
    wf = client.get(f"/qms/documents/{did}/workflow").get_json()
    review_step = next(s for s in wf["steps"] if s["step_key"] == "under_review")
    assert {a["user_id"] for a in review_step["approvers"]} == {"rev-42"}
    assert review_step["approvers"][0]["pool_role"] == "reviewer"


def test_start_workflow_leaves_review_step_unassigned_when_reviewer_not_configured(client):
    """Backward compatibility: a company that hasn't configured a Reviewer
    pool role is completely unaffected — same as before this change."""
    did = _create_submittable(client)
    r = client.post(f"/qms/documents/{did}/workflow/start")
    assert r.status_code == 201, r.get_json()
    wf = client.get(f"/qms/documents/{did}/workflow").get_json()
    review_step = next(s for s in wf["steps"] if s["step_key"] == "under_review")
    assert review_step["approvers"] == []


def test_get_workflow_attaches_pool_summary_to_under_review_only_for_reviewer(client):
    client.post("/qms/documents/approver-pool",
                json={"pool_role": "reviewer", "user_id": "rev-1", "display_name": "Rita"})
    did = _create_submittable(client)
    client.post(f"/qms/documents/{did}/workflow/start")
    wf = client.get(f"/qms/documents/{did}/workflow").get_json()
    review_step = next(s for s in wf["steps"] if s["step_key"] == "under_review")
    assert review_step["pool_summary"] == [{"pool_role": "reviewer", "configured": True, "optional": False}]


def test_get_workflow_effective_step_pool_summary_unaffected_by_reviewer_addition(client):
    """The 'effective' step's pool_summary must still be exactly Department
    Head / Quality Head / Plant Head — adding 'reviewer' to ALL_POOL_ROLES
    must not leak it into this step's summary."""
    client.post("/qms/documents/approver-pool",
                json={"pool_role": "department_head", "user_id": "dh-1", "display_name": "Dana"})
    client.post("/qms/documents/approver-pool",
                json={"pool_role": "quality_head", "user_id": "qh-1", "display_name": "Quinn"})
    did = _create_submittable(client)
    client.post(f"/qms/documents/{did}/workflow/start")
    wf = client.get(f"/qms/documents/{did}/workflow").get_json()
    approval_step = next(s for s in wf["steps"] if s["step_key"] == "effective")
    roles = {p["pool_role"] for p in approval_step["pool_summary"]}
    assert roles == {"department_head", "quality_head", "plant_head"}


def test_reviewer_cannot_decide_step_they_are_not_assigned_to(client, monkeypatch):
    from pharmagpt.auth.context import TenantContext
    import tests.conftest as conftest_module

    client.post("/qms/documents/approver-pool",
                json={"pool_role": "reviewer", "user_id": "rev-assigned", "display_name": "Assigned Reviewer"})
    did = _create_submittable(client)
    client.post(f"/qms/documents/{did}/workflow/start")

    someone_else = TenantContext(
        user_id="not-the-reviewer", email="other@example.com", display_name="Other",
        role="reviewer_qa", company_id=conftest_module._TEST_TENANT.company_id,
    )
    monkeypatch.setattr(conftest_module, "_TEST_TENANT", someone_else)
    r = client.post(f"/qms/documents/{did}/workflow/steps/2/decide", json={"decision": "approve"})
    assert r.status_code in (403, 409)


def test_approver_pool_is_tenant_isolated(client, monkeypatch):
    from pharmagpt.auth.context import TenantContext
    import tests.conftest as conftest_module

    client.post("/qms/documents/approver-pool",
                json={"pool_role": "reviewer", "user_id": "rev-company-a", "display_name": "Company A Reviewer"})

    other_company = TenantContext(
        user_id="admin-b", email="admin-b@example.com", display_name="Admin B",
        role="company_admin", company_id="other-company-xyz",
    )
    monkeypatch.setattr(conftest_module, "_TEST_TENANT", other_company)
    pool_b = client.get("/qms/documents/approver-pool").get_json()
    assert "rev-company-a" not in {p["user_id"] for p in pool_b}
