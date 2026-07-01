"""
tests/test_pdf_engines.py — Unit tests for each extraction engine adapter in
isolation (services/extraction/pdf_engines.py, simple_engines.py).
"""

import pytest

from pharmagpt.services.extraction.base import EngineOpenError, PageExtractionError
from pharmagpt.services.extraction.pdf_engines import (
    OCRPlaceholderEngine,
    PdfplumberEngine,
    PyMuPDFEngine,
    PyPDFEngine,
)
from pharmagpt.services.extraction.simple_engines import DocxEngine, ExcelEngine, TxtEngine

PDF_ENGINES = [PyPDFEngine(), PdfplumberEngine(), PyMuPDFEngine()]


@pytest.mark.parametrize("engine", PDF_ENGINES, ids=lambda e: e.name)
def test_open_and_extract_small_pdf(engine, fixtures_dir):
    handle = engine.open(fixtures_dir["small.pdf"])
    try:
        assert engine.page_count(handle) == 5
        text = engine.extract_page(handle, 0)
        assert text and "Page 1" in text
    finally:
        engine.close(handle)


@pytest.mark.parametrize("engine", PDF_ENGINES, ids=lambda e: e.name)
def test_scanned_pdf_returns_no_text_but_does_not_raise(engine, fixtures_dir):
    handle = engine.open(fixtures_dir["scanned.pdf"])
    try:
        text = engine.extract_page(handle, 0)
        assert not text or not text.strip()
    finally:
        engine.close(handle)


@pytest.mark.parametrize("engine", PDF_ENGINES, ids=lambda e: e.name)
def test_corrupted_pdf_raises_engine_open_error(engine, fixtures_dir):
    with pytest.raises(EngineOpenError):
        engine.open(fixtures_dir["corrupted.pdf"])


@pytest.mark.parametrize("engine", PDF_ENGINES, ids=lambda e: e.name)
def test_password_protected_pdf_raises_engine_open_error(engine, fixtures_dir):
    with pytest.raises(EngineOpenError):
        engine.open(fixtures_dir["password_protected.pdf"])


def test_ocr_placeholder_always_unavailable():
    engine = OCRPlaceholderEngine()
    with pytest.raises(EngineOpenError):
        engine.open("anything.pdf")
    with pytest.raises(PageExtractionError):
        engine.extract_page(None, 0)


def test_docx_engine_single_unit_loop(tmp_path):
    from docx import Document

    path = str(tmp_path / "sample.docx")
    doc = Document()
    doc.add_paragraph("Hello from a DOCX test fixture.")
    doc.add_paragraph("Second paragraph with more words to raise the page estimate substantially. " * 20)
    doc.save(path)

    engine = DocxEngine()
    handle = engine.open(path)
    # Loop unit count must always be 1 — extract_page must not be called
    # more than once per document (that would duplicate the extracted text).
    assert engine.page_count(handle) == 1
    assert engine.display_page_count(handle) >= 1
    text = engine.extract_page(handle, 0)
    assert "Hello from a DOCX test fixture." in text


def test_excel_engine_single_unit_loop(tmp_path):
    import openpyxl

    path = str(tmp_path / "sample.xlsx")
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append(["Parameter", "Result"])
    ws.append(["Temperature", "22.4 C"])
    wb.save(path)

    engine = ExcelEngine()
    handle = engine.open(path)
    assert engine.page_count(handle) == 1
    text = engine.extract_page(handle, 0)
    assert "Temperature" in text


def test_txt_engine_single_unit_loop(tmp_path):
    path = str(tmp_path / "sample.txt")
    with open(path, "w", encoding="utf-8") as fh:
        fh.write("plain text content " * 50)

    engine = TxtEngine()
    handle = engine.open(path)
    assert engine.page_count(handle) == 1
    text = engine.extract_page(handle, 0)
    assert "plain text content" in text
