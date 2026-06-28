"""
services/docx_generator.py — Professional DOCX generation engine for PharmaGPT v0.9.4.

This module is the single source of truth for all Word document formatting.
Every document type (IQ, OQ, PQ, URS, FMEA, CAPA, Deviation, Change Control, etc.)
uses this engine so all exports share the same regulated-industry styling.

Architecture
------------
  PharmaTheme     — centralised colour / font / spacing constants
  DocumentData    — structured input dataclass (no AI, no DB calls here)
  DocxGenerator   — builds the DOCX section-by-section
  generate_docx() — public one-call API → returns bytes

Document structure produced
---------------------------
  Cover Page  →  TOC  →  Revision History  →  Approval Page
  →  Main Body  →  Annexures A–D

Header (every page except cover)
  Logo left  |  Document Title centre  |  Protocol Number right

Footer (every page except cover)
  Controlled Document  |  Revision  |  Page X of Y  |  PharmaGPT  |  Date

Optional
  DRAFT diagonal watermark when document_status == "Draft"

PDF integration point
  The DocxGenerator.build() returns bytes.  A future pdf_generator.py
  can accept the same DocumentData and produce PDF without touching this module.
"""

from __future__ import annotations

import logging
import os
import re
from dataclasses import dataclass, field
from datetime import date, datetime
from io import BytesIO
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

logger = logging.getLogger(__name__)

# ── Output directory ──────────────────────────────────────────────────────────
_GENERATED_DIR = Path(__file__).resolve().parents[3] / "generated_documents"
_GENERATED_DIR.mkdir(parents=True, exist_ok=True)


# ══════════════════════════════════════════════════════════════════════════════
# PharmaTheme — single source of truth for all visual constants
# ══════════════════════════════════════════════════════════════════════════════

class PharmaTheme:
    # Colours
    NAVY        = RGBColor(0x0A, 0x23, 0x42)   # #0A2342 — primary brand
    BLUE        = RGBColor(0x15, 0x65, 0xC0)   # #1565C0 — secondary
    DARK_BLUE   = RGBColor(0x0D, 0x3B, 0x7A)   # #0D3B7A — heading 3
    ACCENT      = RGBColor(0x00, 0xAC, 0xC1)   # #00ACC1 — accent
    WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
    BLACK       = RGBColor(0x1A, 0x2B, 0x3C)   # near-black body text
    GRAY        = RGBColor(0x60, 0x70, 0x80)   # secondary text / borders
    LIGHT_GRAY  = RGBColor(0xF4, 0xF7, 0xFB)  # alternate table rows
    MID_GRAY    = RGBColor(0xD0, 0xD9, 0xE4)  # divider lines

    # Hex strings for XML shading (no # prefix)
    HEX_NAVY        = "0A2342"
    HEX_BLUE        = "1565C0"
    HEX_LIGHT_GRAY  = "F4F7FB"
    HEX_MID_GRAY    = "D0D9E4"
    HEX_ACCENT      = "E8F4FD"   # pale blue for cover metadata rows

    # Fonts
    FONT_BODY = "Calibri"
    FONT_MONO = "Courier New"

    # Point sizes
    SIZE_COVER_COMPANY  = Pt(22)
    SIZE_COVER_SITE     = Pt(14)
    SIZE_COVER_DOCTYPE  = Pt(11)
    SIZE_COVER_TITLE    = Pt(22)
    SIZE_COVER_EQUIP    = Pt(13)
    SIZE_COVER_META     = Pt(10)
    SIZE_H1             = Pt(16)
    SIZE_H2             = Pt(14)
    SIZE_H3             = Pt(12)
    SIZE_BODY           = Pt(11)
    SIZE_TABLE          = Pt(10)
    SIZE_SMALL          = Pt(9)
    SIZE_CAPTION        = Pt(8)

    # Spacing
    SPACE_H1_BEFORE = Pt(18)
    SPACE_H1_AFTER  = Pt(6)
    SPACE_H2_BEFORE = Pt(14)
    SPACE_H2_AFTER  = Pt(4)
    SPACE_H3_BEFORE = Pt(10)
    SPACE_H3_AFTER  = Pt(4)
    SPACE_PARA_AFTER = Pt(6)

    LINE_SPACING = 1.15

    # Page layout
    MARGIN_TOP    = Inches(1.0)
    MARGIN_BOTTOM = Inches(1.0)
    MARGIN_LEFT   = Inches(1.0)
    MARGIN_RIGHT  = Inches(1.0)
    PAGE_W = Cm(21.0)
    PAGE_H = Cm(29.7)


# ══════════════════════════════════════════════════════════════════════════════
# DocumentData — structured input; no AI / DB calls inside
# ══════════════════════════════════════════════════════════════════════════════

@dataclass
class RevisionEntry:
    revision:    str = ""
    date:        str = ""
    description: str = ""
    prepared_by: str = ""
    approved_by: str = ""


@dataclass
class Signatory:
    role:        str = ""
    name:        str = ""
    designation: str = ""


@dataclass
class DocumentData:
    # ── Core ──────────────────────────────────────────────────────────────────
    doc_type:         str = "DOC"
    doc_title:        str = ""
    markdown_content: str = ""

    # ── Project ───────────────────────────────────────────────────────────────
    project_name: str = ""
    project_id:   int = 0

    # ── Equipment ─────────────────────────────────────────────────────────────
    equipment_name: str = ""
    equipment_id:   str = ""
    manufacturer:   str = ""
    model:          str = ""

    # ── Document metadata ─────────────────────────────────────────────────────
    protocol_number:  str = ""
    revision_number:  str = "00"
    effective_date:   str = ""
    document_status:  str = "Draft"   # "Draft" | "Final"

    # ── Company ───────────────────────────────────────────────────────────────
    company_name: str = ""
    site_name:    str = ""
    department:   str = ""

    # ── Personnel ─────────────────────────────────────────────────────────────
    prepared_by: str = ""
    reviewed_by: str = ""
    approved_by: str = ""

    # ── Revision history (list[RevisionEntry | dict]) ─────────────────────────
    revision_history: list = field(default_factory=list)

    # ── Approval-page signatories (list[Signatory | dict]) ───────────────────
    signatories: list = field(default_factory=list)

    # ── Optional ──────────────────────────────────────────────────────────────
    logo_path:  Optional[str] = None
    keywords:   list          = field(default_factory=list)

    # ── Auto-filled ───────────────────────────────────────────────────────────
    generated_date: str = field(
        default_factory=lambda: date.today().strftime("%d-%b-%Y")
    )


# ══════════════════════════════════════════════════════════════════════════════
# DocxGenerator
# ══════════════════════════════════════════════════════════════════════════════

class DocxGenerator:
    """
    Builds a professional pharmaceutical DOCX from a DocumentData instance.

    Usage
    -----
        gen   = DocxGenerator(document_data)
        bytes = gen.build()          # returns raw DOCX bytes
        gen.save("filename.docx")    # saves to /generated_documents/
    """

    T = PharmaTheme   # alias for brevity inside methods

    def __init__(self, data: DocumentData, review_result=None) -> None:
        self.data          = data
        self.doc           = Document()
        self._in_code_block = False
        self._review_result = review_result  # Optional ReviewResult from review engine

    # ── Public API ────────────────────────────────────────────────────────────

    def build(self) -> bytes:
        """Assemble all sections and return DOCX as bytes."""
        self._configure_page_layout()
        self._configure_styles()
        self._set_document_properties()
        self._add_cover_page()
        self._add_page_break()
        self._add_toc_section()
        self._add_page_break()
        self._add_revision_history()
        self._add_page_break()
        self._add_approval_page()
        self._add_page_break()
        self._parse_and_add_content()
        self._add_page_break()
        self._add_annexures()
        if self._review_result is not None:
            self._add_page_break()
            self._add_review_report(self._review_result)
        self._setup_header_footer()
        if self.data.document_status.lower() == "draft":
            self._add_draft_watermark()

        buf = BytesIO()
        self.doc.save(buf)
        return buf.getvalue()

    def save(self, filename: str | None = None) -> Path:
        """Save to /generated_documents/ and return the file path."""
        if not filename:
            safe = re.sub(r"[^\w\-]", "_", self.data.doc_title or self.data.doc_type)
            filename = f"{safe}_{self.data.generated_date}.docx"
        path = _GENERATED_DIR / filename
        path.write_bytes(self.build())
        logger.info("Saved DOCX: %s", path)
        return path

    # ── Page layout & styles ──────────────────────────────────────────────────

    def _configure_page_layout(self) -> None:
        T = self.T
        for section in self.doc.sections:
            section.page_width    = T.PAGE_W
            section.page_height   = T.PAGE_H
            section.top_margin    = T.MARGIN_TOP
            section.bottom_margin = T.MARGIN_BOTTOM
            section.left_margin   = T.MARGIN_LEFT
            section.right_margin  = T.MARGIN_RIGHT
            # Suppress header/footer on the cover page (first page)
            section.different_first_page_header_footer = True

    def _configure_styles(self) -> None:
        T  = self.T
        st = self.doc.styles

        def _set(style_name, size, bold=False, color=None,
                 space_before=Pt(0), space_after=Pt(6)):
            try:
                s = st[style_name]
            except KeyError:
                return
            s.font.name  = T.FONT_BODY
            s.font.size  = size
            s.font.bold  = bold
            if color:
                s.font.color.rgb = color
            s.paragraph_format.space_before       = space_before
            s.paragraph_format.space_after        = space_after
            s.paragraph_format.line_spacing_rule  = WD_LINE_SPACING.MULTIPLE
            s.paragraph_format.line_spacing       = T.LINE_SPACING

        _set("Normal",    T.SIZE_BODY,  color=T.BLACK)
        _set("Heading 1", T.SIZE_H1, bold=True, color=T.NAVY,
             space_before=T.SPACE_H1_BEFORE, space_after=T.SPACE_H1_AFTER)
        _set("Heading 2", T.SIZE_H2, bold=True, color=T.BLUE,
             space_before=T.SPACE_H2_BEFORE, space_after=T.SPACE_H2_AFTER)
        _set("Heading 3", T.SIZE_H3, bold=True, color=T.DARK_BLUE,
             space_before=T.SPACE_H3_BEFORE, space_after=T.SPACE_H3_AFTER)

    def _set_document_properties(self) -> None:
        d = self.data
        cp = self.doc.core_properties
        cp.title    = d.doc_title or f"{d.doc_type} Protocol"
        cp.author   = d.prepared_by or "PharmaGPT"
        cp.company  = d.company_name or ""
        cp.subject  = f"{d.doc_type} — {d.equipment_name}"
        cp.keywords = ", ".join(d.keywords) if d.keywords else d.doc_type
        try:
            cp.revision = max(1, int(d.revision_number or 1))
        except (ValueError, TypeError):
            cp.revision = 1
        cp.created  = datetime.now()
        cp.comments = f"Generated by PharmaGPT on {d.generated_date}"

    # ── Cover Page ────────────────────────────────────────────────────────────

    def _add_cover_page(self) -> None:
        d = self.data
        T = self.T
        doc = self.doc

        def _cp(text="", size=Pt(11), bold=False, color=None,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                space_before=Pt(0), space_after=Pt(6)) -> None:
            p = doc.add_paragraph()
            p.alignment = align
            p.paragraph_format.space_before = space_before
            p.paragraph_format.space_after  = space_after
            if text:
                run = p.add_run(text)
                run.font.name  = T.FONT_BODY
                run.font.size  = size
                run.font.bold  = bold
                run.font.color.rgb = color or T.BLACK

        # ── Logo ─────────────────────────────────────────────────────────────
        if d.logo_path and os.path.isfile(d.logo_path):
            try:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after  = Pt(8)
                run = p.add_run()
                run.add_picture(d.logo_path, height=Inches(1.2))
            except Exception:
                _cp("[Company Logo]", size=T.SIZE_SMALL, color=T.GRAY)
        else:
            _cp("[Company Logo]", size=T.SIZE_SMALL, color=T.GRAY,
                space_before=Pt(20), space_after=Pt(8))

        # ── Company / Site ────────────────────────────────────────────────────
        _cp(d.company_name or "Company Name",
            size=T.SIZE_COVER_COMPANY, bold=True, color=T.NAVY, space_after=Pt(4))
        _cp(d.site_name or "Site / Location",
            size=T.SIZE_COVER_SITE, color=T.BLUE, space_after=Pt(2))
        if d.department:
            _cp(d.department, size=T.SIZE_COVER_DOCTYPE, color=T.GRAY, space_after=Pt(4))

        # ── Divider ───────────────────────────────────────────────────────────
        self._add_cover_rule()

        # ── Document type badge ───────────────────────────────────────────────
        _cp(_DOC_TYPE_LABELS.get(d.doc_type, d.doc_type).upper(),
            size=T.SIZE_COVER_DOCTYPE, color=T.ACCENT, space_before=Pt(10), space_after=Pt(6))

        # ── Document title ────────────────────────────────────────────────────
        _cp(d.doc_title or f"{d.doc_type} Protocol",
            size=T.SIZE_COVER_TITLE, bold=True, color=T.NAVY,
            space_before=Pt(4), space_after=Pt(8))

        # ── Equipment line ────────────────────────────────────────────────────
        equip_parts = [p for p in [d.equipment_name, d.equipment_id] if p]
        _cp(" | ".join(equip_parts) if equip_parts else "Equipment",
            size=T.SIZE_COVER_EQUIP, bold=True, color=T.DARK_BLUE, space_after=Pt(16))

        # ── Metadata table ────────────────────────────────────────────────────
        meta_rows = [
            ("Protocol Number",  d.protocol_number  or "—"),
            ("Revision",         d.revision_number  or "00"),
            ("Effective Date",   d.effective_date   or d.generated_date),
            ("Document Status",  d.document_status  or "Draft"),
            ("Manufacturer",     d.manufacturer     or "—"),
            ("Model",            d.model            or "—"),
        ]
        self._add_cover_meta_table(meta_rows)

        doc.add_paragraph().paragraph_format.space_after = Pt(12)

        # ── Signature stub table ──────────────────────────────────────────────
        sig_rows = [
            ("Prepared By",  d.prepared_by  or ""),
            ("Reviewed By",  d.reviewed_by  or ""),
            ("Approved By",  d.approved_by  or ""),
        ]
        self._add_cover_sig_table(sig_rows)

        # ── Confidential footer ───────────────────────────────────────────────
        doc.add_paragraph().paragraph_format.space_after = Pt(20)
        self._add_cover_rule()
        _cp("CONFIDENTIAL — CONTROLLED DOCUMENT",
            size=T.SIZE_SMALL, bold=True, color=T.GRAY, space_before=Pt(6), space_after=Pt(2))
        company_display = d.company_name or "the Company"
        _cp(
            f"This document is the property of {company_display}. "
            "Unauthorised reproduction or distribution is strictly prohibited.",
            size=Pt(8), color=T.GRAY, space_after=Pt(2),
        )
        _cp(f"Generated by PharmaGPT  ·  {d.generated_date}",
            size=Pt(8), color=T.GRAY, space_after=Pt(0))

    def _add_cover_rule(self) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "12")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), self.T.HEX_NAVY)
        pBdr.append(bot)
        pPr.append(pBdr)

    def _add_cover_meta_table(self, rows: list[tuple[str, str]]) -> None:
        T   = self.T
        tbl = self.doc.add_table(rows=len(rows), cols=2)
        tbl.alignment  = WD_TABLE_ALIGNMENT.CENTER
        tbl.style      = "Table Grid"
        _set_table_width(tbl, Inches(5.0))

        for i, (label, value) in enumerate(rows):
            row = tbl.rows[i]
            # Label cell
            lc = row.cells[0]
            lc.text = ""
            lp = lc.paragraphs[0]
            lr = lp.add_run(label)
            lr.font.name = T.FONT_BODY
            lr.font.size = T.SIZE_COVER_META
            lr.font.bold = True
            lr.font.color.rgb = T.WHITE
            _set_cell_bg(lc, T.HEX_NAVY)
            lc.width = Inches(2.0)

            # Value cell
            vc = row.cells[1]
            vc.text = ""
            vp = vc.paragraphs[0]
            vr = vp.add_run(value)
            vr.font.name = T.FONT_BODY
            vr.font.size = T.SIZE_COVER_META
            vr.font.color.rgb = T.NAVY
            _set_cell_bg(vc, T.HEX_ACCENT)

        self.doc.add_paragraph()

    def _add_cover_sig_table(self, rows: list[tuple[str, str]]) -> None:
        T   = self.T
        tbl = self.doc.add_table(rows=len(rows), cols=3)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style     = "Table Grid"
        _set_table_width(tbl, Inches(5.5))

        headers = ["Role", "Name", "Signature / Date"]
        hrow = tbl.rows[0]   # reuse first data row as header — insert a header row
        # Actually insert a proper header row first
        tbl._tbl.remove(hrow._tr)  # remove, rebuild with header + data

        # Rebuild: header + data rows
        tbl2 = self.doc.add_table(rows=len(rows) + 1, cols=3)
        tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl2.style     = "Table Grid"
        _set_table_width(tbl2, Inches(5.5))

        # Remove the placeholder table we just created
        tbl._tbl.getparent().remove(tbl._tbl)

        # Header row
        hr = tbl2.rows[0]
        for ci, hdr in enumerate(headers):
            c = hr.cells[ci]
            c.text = ""
            rn = c.paragraphs[0].add_run(hdr)
            rn.font.name = T.FONT_BODY
            rn.font.size = T.SIZE_TABLE
            rn.font.bold = True
            rn.font.color.rgb = T.WHITE
            _set_cell_bg(c, T.HEX_NAVY)

        for ri, (role, name) in enumerate(rows):
            row = tbl2.rows[ri + 1]
            for ci, val in enumerate([role, name, ""]):
                c = row.cells[ci]
                c.text = ""
                rn = c.paragraphs[0].add_run(val)
                rn.font.name = T.FONT_BODY
                rn.font.size = T.SIZE_TABLE
                rn.font.color.rgb = T.BLACK

        self.doc.add_paragraph()

    # ── Table of Contents ─────────────────────────────────────────────────────

    def _add_toc_section(self) -> None:
        T   = self.T
        doc = self.doc

        # Section heading (not a Word Heading style — excluded from TOC itself)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("TABLE OF CONTENTS")
        run.font.name  = T.FONT_BODY
        run.font.size  = T.SIZE_H2
        run.font.bold  = True
        run.font.color.rgb = T.NAVY
        p.paragraph_format.space_after = Pt(16)
        self._add_blue_underline(p)

        # TOC field — Word refreshes this when the file is opened
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(4)
        run  = para.add_run()
        _add_field(run, 'TOC \\o "1-3" \\h \\z \\u')

        # Instruction note
        note = doc.add_paragraph()
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        nr = note.add_run(
            "[ Right-click the table above and select 'Update Field' to refresh page numbers ]"
        )
        nr.font.name      = T.FONT_BODY
        nr.font.size      = Pt(8)
        nr.font.italic    = True
        nr.font.color.rgb = T.GRAY
        note.paragraph_format.space_before = Pt(6)

    # ── Revision History ──────────────────────────────────────────────────────

    def _add_revision_history(self) -> None:
        T   = self.T
        doc = self.doc
        d   = self.data

        self._add_section_heading("REVISION HISTORY")

        headers = ["Revision", "Date", "Description of Change", "Prepared By", "Approved By"]
        history = d.revision_history or [
            RevisionEntry(
                revision    = d.revision_number or "00",
                date        = d.generated_date,
                description = "Initial issue",
                prepared_by = d.prepared_by or "",
                approved_by = d.approved_by or "",
            )
        ]

        # Normalise to RevisionEntry
        entries: list[RevisionEntry] = []
        for h in history:
            if isinstance(h, dict):
                entries.append(RevisionEntry(**{k: v for k, v in h.items()
                                                if k in RevisionEntry.__dataclass_fields__}))
            elif isinstance(h, RevisionEntry):
                entries.append(h)

        tbl = doc.add_table(rows=len(entries) + 1, cols=5)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        tbl.style     = "Table Grid"
        _set_table_width(tbl, Inches(6.5))

        # Header row
        hr = tbl.rows[0]
        _set_repeat_header(hr)
        for ci, hdr in enumerate(headers):
            c = hr.cells[ci]
            c.text = ""
            rn = c.paragraphs[0].add_run(hdr)
            rn.font.name      = T.FONT_BODY
            rn.font.size      = T.SIZE_TABLE
            rn.font.bold      = True
            rn.font.color.rgb = T.WHITE
            _set_cell_bg(c, T.HEX_NAVY)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        for ri, entry in enumerate(entries):
            row    = tbl.rows[ri + 1]
            values = [entry.revision, entry.date, entry.description,
                      entry.prepared_by, entry.approved_by]
            bg = T.HEX_LIGHT_GRAY if ri % 2 == 0 else "FFFFFF"
            for ci, val in enumerate(values):
                c = row.cells[ci]
                c.text = ""
                rn = c.paragraphs[0].add_run(str(val))
                rn.font.name      = T.FONT_BODY
                rn.font.size      = T.SIZE_TABLE
                rn.font.color.rgb = T.BLACK
                _set_cell_bg(c, bg)
                c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        doc.add_paragraph()

    # ── Approval Page ─────────────────────────────────────────────────────────

    def _add_approval_page(self) -> None:
        T   = self.T
        doc = self.doc
        d   = self.data

        self._add_section_heading("APPROVAL PAGE")

        intro = doc.add_paragraph(
            "By signing below, the signatories confirm that they have reviewed this document "
            "and that its content is accurate, complete, and fit for its intended purpose."
        )
        intro.paragraph_format.space_after = Pt(12)
        intro.runs[0].font.name  = T.FONT_BODY
        intro.runs[0].font.size  = T.SIZE_BODY
        intro.runs[0].font.color.rgb = T.BLACK

        # Build signatory list
        default_sigs = [
            Signatory("Prepared By",         d.prepared_by, "Validation Engineer"),
            Signatory("Reviewed By",          d.reviewed_by, "Process Engineer"),
            Signatory("Engineering",          "",            "Engineering Manager"),
            Signatory("Production",           "",            "Production Manager"),
            Signatory("Quality Assurance",    "",            "QA Manager"),
            Signatory("Validation",           "",            "Validation Manager"),
            Signatory("Final Approval",       d.approved_by, "Site Director / VP Quality"),
        ]
        sigs = d.signatories or default_sigs
        entries: list[Signatory] = []
        for s in sigs:
            if isinstance(s, dict):
                entries.append(Signatory(**{k: v for k, v in s.items()
                                            if k in Signatory.__dataclass_fields__}))
            else:
                entries.append(s)

        headers = ["Function / Role", "Name", "Designation", "Signature", "Date"]
        tbl = doc.add_table(rows=len(entries) + 1, cols=5)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        tbl.style     = "Table Grid"
        _set_table_width(tbl, Inches(6.5))

        # Header
        hr = tbl.rows[0]
        _set_repeat_header(hr)
        for ci, hdr in enumerate(headers):
            c = hr.cells[ci]
            c.text = ""
            rn = c.paragraphs[0].add_run(hdr)
            rn.font.name      = T.FONT_BODY
            rn.font.size      = T.SIZE_TABLE
            rn.font.bold      = True
            rn.font.color.rgb = T.WHITE
            _set_cell_bg(c, T.HEX_NAVY)

        for ri, sig in enumerate(entries):
            row    = tbl.rows[ri + 1]
            values = [sig.role, sig.name, sig.designation, "", ""]
            bg = T.HEX_LIGHT_GRAY if ri % 2 == 0 else "FFFFFF"
            for ci, val in enumerate(values):
                c = row.cells[ci]
                c.text = ""
                rn = c.paragraphs[0].add_run(val)
                rn.font.name      = T.FONT_BODY
                rn.font.size      = T.SIZE_TABLE
                rn.font.color.rgb = T.BLACK
                if ci == 3:
                    # Signature cell — give extra height hint
                    c.paragraphs[0].paragraph_format.space_before = Pt(14)
                    c.paragraphs[0].paragraph_format.space_after  = Pt(14)
                _set_cell_bg(c, bg)

        doc.add_paragraph()

    # ── Main Content — Markdown parser ────────────────────────────────────────

    def _parse_and_add_content(self) -> None:
        content = self.data.markdown_content
        if not content:
            p = self.doc.add_paragraph("[ No content generated. ]")
            p.runs[0].font.color.rgb = self.T.GRAY
            return

        lines         = content.split("\n")
        table_rows:   list[list[str]] = []
        code_lines:   list[str]       = []
        in_code_block = False

        def flush_table():
            if table_rows:
                self._add_md_table(table_rows)
                table_rows.clear()

        def flush_code():
            if code_lines:
                self._add_code_block(code_lines)
                code_lines.clear()

        i = 0
        while i < len(lines):
            raw     = lines[i]
            stripped = raw.strip()

            # ── Code fence ────────────────────────────────────────────────────
            if stripped.startswith("```"):
                if in_code_block:
                    flush_code()
                    in_code_block = False
                else:
                    flush_table()
                    in_code_block = True
                i += 1
                continue

            if in_code_block:
                code_lines.append(raw)
                i += 1
                continue

            # ── Table ─────────────────────────────────────────────────────────
            if stripped.startswith("|"):
                if re.match(r"^\|[\s\-:]+(\|[\s\-:]+)*\|?$", stripped):
                    i += 1
                    continue
                cells = [c.strip() for c in stripped.split("|")]
                if cells and cells[0]  == "": cells = cells[1:]
                if cells and cells[-1] == "": cells = cells[:-1]
                table_rows.append(cells)
                i += 1
                continue
            else:
                flush_table()

            # ── Headings ──────────────────────────────────────────────────────
            if stripped.startswith("#### "):
                self._add_md_heading(stripped[5:], level=3)
            elif stripped.startswith("### "):
                self._add_md_heading(stripped[4:], level=2)
            elif stripped.startswith("## "):
                self._add_md_heading(stripped[3:], level=1)
            elif stripped.startswith("# "):
                self._add_md_heading(stripped[2:], level=1)

            # ── Lists ─────────────────────────────────────────────────────────
            elif stripped.startswith("- ") or stripped.startswith("* "):
                self._add_md_bullet(stripped[2:])
            elif re.match(r"^\d+\.\s", stripped):
                self._add_md_numbered(re.sub(r"^\d+\.\s+", "", stripped))

            # ── Horizontal rule ───────────────────────────────────────────────
            elif re.match(r"^-{3,}$|^\*{3,}$|^_{3,}$", stripped):
                self._add_hr()

            # ── Blank line ────────────────────────────────────────────────────
            elif not stripped:
                pass

            # ── Normal paragraph ──────────────────────────────────────────────
            else:
                self._add_md_paragraph(stripped)

            i += 1

        flush_table()
        flush_code()

    # ── Markdown element builders ─────────────────────────────────────────────

    def _add_md_heading(self, text: str, level: int) -> None:
        style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
        style = style_map.get(level, "Heading 3")
        p = self.doc.add_paragraph(style=style)
        _apply_inline(p, text, font=self.T.FONT_BODY, size=None)
        if level == 1:
            self._add_blue_underline(p)

    def _add_md_paragraph(self, text: str) -> None:
        T = self.T
        p = self.doc.add_paragraph(style="Normal")
        p.paragraph_format.space_after = T.SPACE_PARA_AFTER
        _apply_inline(p, text, font=T.FONT_BODY, size=T.SIZE_BODY)

    def _add_md_bullet(self, text: str) -> None:
        T = self.T
        try:
            p = self.doc.add_paragraph(style="List Bullet")
        except KeyError:
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(2)
        _apply_inline(p, text, font=T.FONT_BODY, size=T.SIZE_BODY)

    def _add_md_numbered(self, text: str) -> None:
        T = self.T
        try:
            p = self.doc.add_paragraph(style="List Number")
        except KeyError:
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(2)
        _apply_inline(p, text, font=T.FONT_BODY, size=T.SIZE_BODY)

    def _add_md_table(self, rows: list[list[str]]) -> None:
        T = self.T
        if not rows:
            return
        col_count = max(len(r) for r in rows)
        rows = [r + [""] * (col_count - len(r)) for r in rows]

        tbl = self.doc.add_table(rows=len(rows), cols=col_count)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        tbl.style     = "Table Grid"

        for r_idx, row_data in enumerate(rows):
            row = tbl.rows[r_idx]
            if r_idx == 0:
                _set_repeat_header(row)
            for c_idx, cell_text in enumerate(row_data):
                cell = row.cells[c_idx]
                cell.text = ""
                p  = cell.paragraphs[0]
                rn = p.add_run(cell_text)
                rn.font.name = T.FONT_BODY
                rn.font.size = T.SIZE_TABLE
                if r_idx == 0:
                    rn.bold = True
                    rn.font.color.rgb = T.WHITE
                    _set_cell_bg(cell, T.HEX_NAVY)
                else:
                    rn.font.color.rgb = T.BLACK
                    _set_cell_bg(cell, T.HEX_LIGHT_GRAY if r_idx % 2 == 0 else "FFFFFF")
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        self.doc.add_paragraph()

    def _add_code_block(self, lines: list[str]) -> None:
        T = self.T
        for line in lines:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            p.paragraph_format.left_indent  = Inches(0.25)
            rn = p.add_run(line)
            rn.font.name = T.FONT_MONO
            rn.font.size = Pt(9)
            rn.font.color.rgb = T.NAVY
            _set_para_shading(p, "EEF2F7")
        self.doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def _add_hr(self) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "6")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), self.T.HEX_MID_GRAY)
        pBdr.append(bot)
        pPr.append(pBdr)

    # ── Annexures ─────────────────────────────────────────────────────────────

    def _add_annexures(self) -> None:
        annexures = [
            ("Annexure A", "Equipment Drawings",
             "Insert equipment layout drawings, P&ID diagrams, and dimensional drawings."),
            ("Annexure B", "Calibration Certificates",
             "Attach current calibration certificates for all instruments used during qualification."),
            ("Annexure C", "Utilities Verification",
             "Attach utility qualification data: compressed air, purified water, nitrogen, HVAC."),
            ("Annexure D", "Training Records",
             "Attach training records for all personnel involved in execution of this protocol."),
        ]
        for code, title, placeholder in annexures:
            self._add_section_heading(f"{code} — {title}")
            p = self.doc.add_paragraph()
            rn = p.add_run(placeholder)
            rn.font.name      = self.T.FONT_BODY
            rn.font.size      = self.T.SIZE_BODY
            rn.font.italic    = True
            rn.font.color.rgb = self.T.GRAY
            p.paragraph_format.space_after = Pt(6)

            placeholder_p = self.doc.add_paragraph()
            placeholder_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pr = placeholder_p.add_run(f"[ {title} — attach here ]")
            pr.font.name      = self.T.FONT_BODY
            pr.font.size      = self.T.SIZE_BODY
            pr.font.italic    = True
            pr.font.color.rgb = self.T.GRAY
            placeholder_p.paragraph_format.space_before = Pt(20)
            placeholder_p.paragraph_format.space_after  = Pt(20)
            self._add_hr()
            self.doc.add_paragraph()

    # ── Validation Review Report appendix ─────────────────────────────────────

    def _add_review_report(self, review_result) -> None:
        """Append the Validation Review Report as the final document section."""
        from review.review_formatter import format_docx_sections

        T    = self.T
        data = format_docx_sections(review_result)

        # ── Section heading ───────────────────────────────────────────────────
        h = self.doc.add_heading("Validation Review Report", level=1)
        h.style = self.doc.styles["Heading 1"]
        self.doc.add_paragraph()

        # ── Score banner ──────────────────────────────────────────────────────
        banner_tbl = self.doc.add_table(rows=1, cols=2)
        banner_tbl.style = "Table Grid"
        _remove_table_borders(banner_tbl)
        banner_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

        score_cell = banner_tbl.rows[0].cells[0]
        _set_cell_bg(score_cell, T.HEX_NAVY)
        sp = score_cell.paragraphs[0]
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = sp.add_run(f"{data['score']:.1f} / 100")
        sr.font.name      = T.FONT_BODY
        sr.font.size      = Pt(28)
        sr.font.bold      = True
        sr.font.color.rgb = T.WHITE
        sp.paragraph_format.space_before = Pt(8)
        sp.paragraph_format.space_after  = Pt(4)

        sub_p = score_cell.add_paragraph("Overall Validation Score")
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_r = sub_p.add_run("")
        sub_p.runs[0].font.name      = T.FONT_BODY
        sub_p.runs[0].font.size      = Pt(9)
        sub_p.runs[0].font.color.rgb = T.WHITE

        ready_cell = banner_tbl.rows[0].cells[1]
        hex_ready  = data["readiness_hex"]
        _set_cell_bg(ready_cell, hex_ready)
        rp = ready_cell.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr_run = rp.add_run(data["readiness"])
        rr_run.font.name      = T.FONT_BODY
        rr_run.font.size      = Pt(14)
        rr_run.font.bold      = True
        rr_run.font.color.rgb = T.WHITE
        rp.paragraph_format.space_before = Pt(14)
        rp.paragraph_format.space_after  = Pt(8)

        self.doc.add_paragraph()

        # ── Category scores table ─────────────────────────────────────────────
        h2 = self.doc.add_heading("Category Scores", level=2)
        h2.style = self.doc.styles["Heading 2"]

        cat_tbl = self.doc.add_table(rows=1, cols=3)
        cat_tbl.style = "Table Grid"
        cat_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

        for i, cell in enumerate(cat_tbl.rows[0].cells):
            _set_cell_bg(cell, T.HEX_NAVY)
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(["Category", "Score", "Max"][i])
            r.font.name      = T.FONT_BODY
            r.font.size      = T.SIZE_TABLE
            r.font.bold      = True
            r.font.color.rgb = T.WHITE
            p.alignment      = WD_ALIGN_PARAGRAPH.CENTER

        for idx, (label, score, max_score) in enumerate(data["category_rows"]):
            row = cat_tbl.add_row()
            bg  = T.HEX_LIGHT_GRAY if idx % 2 == 0 else "FFFFFF"
            for cell in row.cells:
                _set_cell_bg(cell, bg)

            row.cells[0].text = ""
            _apply_inline(row.cells[0].paragraphs[0], label, T.FONT_BODY, T.SIZE_TABLE)
            row.cells[1].text = ""
            row.cells[1].paragraphs[0].add_run(f"{score:.1f}").font.size = T.SIZE_TABLE
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row.cells[2].text = ""
            row.cells[2].paragraphs[0].add_run(f"{max_score:.0f}").font.size = T.SIZE_TABLE
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_paragraph()

        # ── Compliance matrix ─────────────────────────────────────────────────
        h2b = self.doc.add_heading("Compliance Matrix", level=2)
        h2b.style = self.doc.styles["Heading 2"]

        cmp_tbl = self.doc.add_table(rows=1, cols=3)
        cmp_tbl.style = "Table Grid"
        cmp_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

        for i, cell in enumerate(cmp_tbl.rows[0].cells):
            _set_cell_bg(cell, T.HEX_NAVY)
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(["Regulation / Framework", "Status", "Notes"][i])
            r.font.name      = T.FONT_BODY
            r.font.size      = T.SIZE_TABLE
            r.font.bold      = True
            r.font.color.rgb = T.WHITE
            p.alignment      = WD_ALIGN_PARAGRAPH.CENTER

        for idx, (regulation, status, colour, explanation) in enumerate(data["compliance"]):
            row = cmp_tbl.add_row()
            bg  = T.HEX_LIGHT_GRAY if idx % 2 == 0 else "FFFFFF"
            _set_cell_bg(row.cells[0], bg)
            _set_cell_bg(row.cells[2], bg)

            row.cells[0].text = ""
            row.cells[0].paragraphs[0].add_run(regulation).font.size = T.SIZE_TABLE
            _set_cell_bg(row.cells[1], colour)
            row.cells[1].text = ""
            sr = row.cells[1].paragraphs[0].add_run(status)
            sr.font.size      = T.SIZE_TABLE
            sr.font.bold      = True
            sr.font.color.rgb = T.WHITE
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row.cells[2].text = ""
            row.cells[2].paragraphs[0].add_run(explanation).font.size = T.SIZE_SMALL

        self.doc.add_paragraph()

        # ── Issue summary table ───────────────────────────────────────────────
        if data["issues"]:
            h2c = self.doc.add_heading("Issue Summary", level=2)
            h2c.style = self.doc.styles["Heading 2"]

            iss_tbl = self.doc.add_table(rows=1, cols=5)
            iss_tbl.style = "Table Grid"
            iss_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

            for i, cell in enumerate(iss_tbl.rows[0].cells):
                _set_cell_bg(cell, T.HEX_NAVY)
                cell.text = ""
                p = cell.paragraphs[0]
                r = p.add_run(
                    ["Rule ID", "Severity", "Section", "Description", "Recommendation"][i]
                )
                r.font.name      = T.FONT_BODY
                r.font.size      = T.SIZE_TABLE
                r.font.bold      = True
                r.font.color.rgb = T.WHITE
                p.alignment      = WD_ALIGN_PARAGRAPH.CENTER

            for idx, (rule_id, severity, sev_colour, section, desc, rec) in enumerate(data["issues"]):
                row = iss_tbl.add_row()
                bg  = T.HEX_LIGHT_GRAY if idx % 2 == 0 else "FFFFFF"
                for c in row.cells:
                    _set_cell_bg(c, bg)
                _set_cell_bg(row.cells[1], sev_colour)

                row.cells[0].text = rule_id
                row.cells[0].paragraphs[0].runs[0].font.size = T.SIZE_TABLE

                row.cells[1].text = ""
                sr = row.cells[1].paragraphs[0].add_run(severity)
                sr.font.size      = T.SIZE_TABLE
                sr.font.bold      = True
                sr.font.color.rgb = T.WHITE
                row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                row.cells[2].text = section
                row.cells[2].paragraphs[0].runs[0].font.size = T.SIZE_TABLE
                row.cells[3].text = desc
                row.cells[3].paragraphs[0].runs[0].font.size = T.SIZE_SMALL
                row.cells[4].text = rec
                row.cells[4].paragraphs[0].runs[0].font.size = T.SIZE_SMALL

            self.doc.add_paragraph()

        # ── Recommendations ───────────────────────────────────────────────────
        if data["recommendations"]:
            h2d = self.doc.add_heading("Recommendations", level=2)
            h2d.style = self.doc.styles["Heading 2"]
            for rec in data["recommendations"]:
                p = self.doc.add_paragraph(style="List Bullet")
                p.add_run(rec).font.size = T.SIZE_BODY
                p.paragraph_format.space_after = Pt(4)
            self.doc.add_paragraph()

        # ── Reviewer comments ─────────────────────────────────────────────────
        h2e = self.doc.add_heading("Reviewer Comments", level=2)
        h2e.style = self.doc.styles["Heading 2"]
        p = self.doc.add_paragraph()
        p.add_run(data["reviewer_comments"]).font.size = T.SIZE_BODY
        p.paragraph_format.space_after = Pt(10)

        # ── Approval recommendation ───────────────────────────────────────────
        h2f = self.doc.add_heading("Approval Recommendation", level=2)
        h2f.style = self.doc.styles["Heading 2"]
        ap = self.doc.add_paragraph()
        ar = ap.add_run(data["approval_recommendation"])
        ar.font.size = T.SIZE_BODY
        ar.font.bold = True
        ap.paragraph_format.space_after = Pt(12)

        # Disclaimer line
        dp = self.doc.add_paragraph()
        dr = dp.add_run(
            "This review was generated automatically by the PharmaGPT Review Engine v0.9.5. "
            "It is not a substitute for formal QA review by qualified personnel."
        )
        dr.font.size      = T.SIZE_SMALL
        dr.font.italic    = True
        dr.font.color.rgb = T.GRAY

    # ── Header / Footer ───────────────────────────────────────────────────────

    def _setup_header_footer(self) -> None:
        """Apply 3-zone header and detailed footer to every page except the cover."""
        d       = self.T
        section = self.doc.sections[0]

        # ── Regular-page header ───────────────────────────────────────────────
        header = section.header
        for p in header.paragraphs:
            p._element.getparent().remove(p._element)

        htbl = header.add_table(rows=1, cols=3, width=Inches(6.5))
        htbl.style = "Table Grid"
        _remove_table_borders(htbl)
        htbl.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Left: company name (or logo placeholder)
        lc = htbl.rows[0].cells[0]
        lc.text = ""
        lr = lc.paragraphs[0].add_run(self.data.company_name or "Company")
        lr.font.name      = self.T.FONT_BODY
        lr.font.size      = self.T.SIZE_CAPTION
        lr.font.bold      = True
        lr.font.color.rgb = self.T.NAVY
        lc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Centre: document title
        cc = htbl.rows[0].cells[1]
        cc.text = ""
        cr = cc.paragraphs[0].add_run(self.data.doc_title or f"{self.data.doc_type} Protocol")
        cr.font.name      = self.T.FONT_BODY
        cr.font.size      = self.T.SIZE_CAPTION
        cr.font.color.rgb = self.T.BLACK
        cc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Right: protocol number
        rc = htbl.rows[0].cells[2]
        rc.text = ""
        rr = rc.paragraphs[0].add_run(self.data.protocol_number or "")
        rr.font.name      = self.T.FONT_BODY
        rr.font.size      = self.T.SIZE_CAPTION
        rr.font.color.rgb = self.T.GRAY
        rc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Bottom border on header table's paragraph
        _add_paragraph_border_bottom(header.add_paragraph(), self.T.HEX_MID_GRAY)

        # ── Regular-page footer ───────────────────────────────────────────────
        footer = section.footer
        for p in footer.paragraphs:
            p._element.getparent().remove(p._element)

        fp = footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        def _frun(text, bold=False):
            r = fp.add_run(text)
            r.font.name      = self.T.FONT_BODY
            r.font.size      = self.T.SIZE_CAPTION
            r.font.bold      = bold
            r.font.color.rgb = self.T.GRAY
            return r

        _frun("CONTROLLED DOCUMENT", bold=True)
        _frun(f"  |  Rev. {self.data.revision_number}  |  Page ")
        _add_page_field(fp, "PAGE")
        _frun(" of ")
        _add_page_field(fp, "NUMPAGES")
        _frun(f"  |  PharmaGPT  |  {self.data.generated_date}")

        _add_paragraph_border_top(fp, self.T.HEX_MID_GRAY)

        # ── First-page header/footer: leave empty (clean cover page) ──────────
        # (already set by different_first_page_header_footer = True in _configure_page_layout)

    # ── DRAFT watermark ───────────────────────────────────────────────────────

    def _add_draft_watermark(self) -> None:
        """Insert a grey diagonal DRAFT watermark via header VML shape."""
        section = self.doc.sections[0]
        header  = section.header

        watermark_xml = (
            '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:pPr><w:jc w:val="center"/></w:pPr>'
            '<w:r><w:rPr/><w:pict>'
            '<v:shape xmlns:v="urn:schemas-microsoft-com:vml" '
            'xmlns:o="urn:schemas-microsoft-com:office:office" '
            'xmlns:w10="urn:schemas-microsoft-com:office:word" '
            'id="_x0000_s2049" type="#_x0000_t136" '
            'style="position:absolute;margin-left:0;margin-top:0;'
            'width:528pt;height:264pt;z-index:-251654144;'
            'mso-position-horizontal:center;mso-position-vertical:center" '
            'fillcolor="#C0C0C0" stroked="f" coordsize="21600,21600">'
            '<v:fill o:detectmouseclick="t"/>'
            '<v:path o:connecttype="none"/>'
            '<v:textbox style="mso-fit-shape-to-text:t"><w:txbxContent/></v:textbox>'
            '<w10:wrap anchorx="page" anchory="page"/>'
            '<v:textpath style="font-family:\'Calibri\';font-size:1pt;'
            'v-text-kern:t" trim="t" fitshape="t" string="DRAFT"/>'
            '</v:shape></w:pict></w:r></w:p>'
        )
        from lxml import etree
        wm_elem = etree.fromstring(watermark_xml)
        header._element.append(wm_elem)

    # ── Shared helpers ────────────────────────────────────────────────────────

    def _add_section_heading(self, text: str) -> None:
        T = self.T
        p = self.doc.add_paragraph(style="Heading 1")
        for run in p.runs:
            run.clear()
        if p.runs:
            p.runs[0].text = text
        else:
            rn = p.add_run(text)
            rn.font.name      = T.FONT_BODY
            rn.font.size      = T.SIZE_H1
            rn.font.bold      = True
            rn.font.color.rgb = T.NAVY
        p.paragraph_format.space_before = T.SPACE_H1_BEFORE
        p.paragraph_format.space_after  = T.SPACE_H1_AFTER
        self._add_blue_underline(p)

    def _add_blue_underline(self, p) -> None:
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "8")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), self.T.HEX_BLUE)
        pBdr.append(bot)
        pPr.append(pBdr)

    def _add_page_break(self) -> None:
        p   = self.doc.add_paragraph()
        run = p.add_run()
        run.add_break(docx_page_break())


# ══════════════════════════════════════════════════════════════════════════════
# Module-level helpers (pure functions, reusable by pdf_generator.py)
# ══════════════════════════════════════════════════════════════════════════════

def _set_cell_bg(cell, hex_color: str) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_para_shading(p, hex_color: str) -> None:
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    pPr.append(shd)


def _set_table_width(tbl, width) -> None:
    tblPr  = tbl._tbl.tblPr
    tblW   = OxmlElement("w:tblW")
    # Convert to twentieths-of-a-point (twips): Inches(x) is already in EMU
    twips  = int(width.pt * 20)
    tblW.set(qn("w:w"),    str(twips))
    tblW.set(qn("w:type"), "dxa")
    existing = tblPr.find(qn("w:tblW"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblW)


def _set_repeat_header(row) -> None:
    """Mark a table row as a repeating header (appears on each new page)."""
    trPr  = row._tr.get_or_add_trPr()
    tblHdr = OxmlElement("w:tblHeader")
    trPr.append(tblHdr)


def _remove_table_borders(tbl) -> None:
    """Remove all borders from a table (used for header layout table)."""
    tblPr  = tbl._tbl.tblPr
    tblBdr = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "none")
        el.set(qn("w:sz"),    "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tblBdr.append(el)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBdr)


def _add_field(run, field_code: str) -> None:
    """Insert a Word field (TOC, PAGE, NUMPAGES, etc.) into a run."""
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_code} "

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def _add_page_field(p, field_name: str) -> None:
    """Append a PAGE or NUMPAGES field to an existing paragraph."""
    r = p.add_run()
    r.font.name = PharmaTheme.FONT_BODY
    r.font.size = PharmaTheme.SIZE_CAPTION
    r.font.color.rgb = PharmaTheme.GRAY
    _add_field(r, field_name)


def _add_paragraph_border_bottom(p, hex_color: str) -> None:
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), hex_color)
    pBdr.append(bot)
    pPr.append(pBdr)


def _add_paragraph_border_top(p, hex_color: str) -> None:
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top  = OxmlElement("w:top")
    top.set(qn("w:val"),   "single")
    top.set(qn("w:sz"),    "4")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), hex_color)
    pBdr.append(top)
    pPr.append(pBdr)


def _apply_inline(p, text: str, font: str = "Calibri",
                  size: Optional[Pt] = None) -> None:
    """
    Split text on **bold**, *italic*, and `code` markers and add styled runs.
    Preserves unformatted text between markers.
    """
    pattern = r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)"
    parts   = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.name = font
            if size: run.font.size = size
        elif part.startswith("*") and part.endswith("*"):
            run = p.add_run(part[1:-1])
            run.italic = True
            run.font.name = font
            if size: run.font.size = size
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = PharmaTheme.FONT_MONO
            if size: run.font.size = size
        else:
            run = p.add_run(part)
            run.font.name = font
            if size: run.font.size = size


def docx_page_break():
    """Return the WD_BREAK enum value for a page break."""
    from docx.enum.text import WD_BREAK
    return WD_BREAK.PAGE


# ── Document-type label map ───────────────────────────────────────────────────
_DOC_TYPE_LABELS = {
    "IQ":             "Installation Qualification Protocol",
    "OQ":             "Operational Qualification Protocol",
    "PQ":             "Performance Qualification Protocol",
    "URS":            "User Requirements Specification",
    "DQ":             "Design Qualification Protocol",
    "FAT":            "Factory Acceptance Test Protocol",
    "SAT":            "Site Acceptance Test Protocol",
    "FMEA":           "Failure Mode and Effects Analysis",
    "CAPA":           "Corrective and Preventive Action Report",
    "Deviation":      "Deviation Report",
    "Change Control": "Change Control Record",
}


# ══════════════════════════════════════════════════════════════════════════════
# Public API
# ══════════════════════════════════════════════════════════════════════════════

def generate_docx(document_data: DocumentData, review_result=None) -> bytes:
    """
    Generate a professional DOCX from a DocumentData instance.

    This is the primary integration point.  Call it from any route or service:

        from services.docx_generator import generate_docx, build_document_data
        docx_bytes = generate_docx(build_document_data(form_data, doc_type, content))

    Parameters
    ----------
    document_data : DocumentData — the structured document input.
    review_result : Optional ReviewResult — if provided, a Validation Review
                    Report is appended as the final section of the DOCX.

    Returns raw DOCX bytes suitable for Flask send_file().

    PDF integration point
    ---------------------
    A future pdf_generator.py should accept the same DocumentData and use a
    headless-Word or WeasyPrint pipeline, calling the same helper functions
    (_set_cell_bg, _apply_inline, etc.) to guarantee formatting parity.
    """
    return DocxGenerator(document_data, review_result=review_result).build()


def build_document_data(
    form_data:  dict,
    doc_type:   str,
    title:      str,
    content:    str,
    project_id: int = 0,
) -> DocumentData:
    """
    Convenience factory: converts the wizard form_data dict into a DocumentData.

    form_data keys recognised (all optional):
        project_name, equipment_name, equipment_id, model, manufacturer,
        protocol_number, revision_number, effective_date, document_status,
        company_name, site_name, department, prepared_by, reviewed_by,
        approved_by, logo_path, details (nested dict)
    """
    fd = form_data or {}

    return DocumentData(
        doc_type         = doc_type,
        doc_title        = title,
        markdown_content = content,
        project_name     = fd.get("project_name", ""),
        project_id       = project_id,
        equipment_name   = fd.get("equipment_name", ""),
        equipment_id     = fd.get("equipment_id", ""),
        manufacturer     = fd.get("manufacturer", ""),
        model            = fd.get("model", ""),
        protocol_number  = fd.get("protocol_number", ""),
        revision_number  = fd.get("revision_number", "00"),
        effective_date   = fd.get("effective_date", date.today().strftime("%d-%b-%Y")),
        document_status  = fd.get("document_status", "Draft"),
        company_name     = fd.get("company_name", ""),
        site_name        = fd.get("site_name", ""),
        department       = fd.get("department", ""),
        prepared_by      = fd.get("prepared_by", ""),
        reviewed_by      = fd.get("reviewed_by", ""),
        approved_by      = fd.get("approved_by", ""),
        logo_path        = fd.get("logo_path"),
        keywords         = [doc_type, fd.get("equipment_name", ""), "Pharmaceutical"],
        revision_history = fd.get("revision_history", []),
        signatories      = fd.get("signatories", []),
    )
