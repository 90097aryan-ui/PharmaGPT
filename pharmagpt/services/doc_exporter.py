"""
services/doc_exporter.py — Convert generated markdown to DOCX.

Handles the DOCX export path. PDF export is done client-side via
window.print() with print-optimised CSS (produces professional output
without requiring server-side system libraries like WeasyPrint/GTK).

DOCX pipeline:
  markdown_to_docx(content, doc_type, form_data)
    → parses markdown line-by-line
    → builds a python-docx Document with pharmaceutical styling
    → returns raw bytes (served as application/vnd.openxmlformats-…)

Supported markdown elements:
  # Title        → Heading level 0 (large, centered, navy)
  ## Section     → Heading level 1 (navy)
  ### Sub-section → Heading level 2 (blue)
  | table |      → Word table with styled header row
  - bullet       → List Bullet style
  1. numbered    → List Number style
  **bold text**  → Bold run inside paragraph
  blank line     → paragraph separator
  --- separator  → thin horizontal line paragraph
  regular text   → Normal paragraph (with inline bold/italic support)
"""

import re
from io import BytesIO
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Color palette (matches PharmaGPT UI) ─────────────────────────────────────
NAVY  = RGBColor(0x0A, 0x23, 0x42)   # #0A2342
BLUE  = RGBColor(0x15, 0x65, 0xC0)   # #1565C0
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x1A, 0x2B, 0x3C)
GRAY  = RGBColor(0x60, 0x70, 0x80)


# ── Public API ────────────────────────────────────────────────────────────────

def markdown_to_docx(content: str, doc_type: str, form_data: dict) -> bytes:
    """
    Convert markdown content (from Gemini) to a DOCX file.

    Parameters
    ----------
    content   : generated markdown text
    doc_type  : "OQ" | "IQ" | etc.   (used for filename / cover styling)
    form_data : wizard form data (for header metadata)

    Returns
    -------
    bytes — the raw DOCX file, ready to serve as a download
    """
    doc = Document()
    _configure_page(doc)
    _add_page_header(doc, doc_type, form_data)
    _parse_markdown(doc, content)
    _add_footer(doc, form_data)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Page layout ───────────────────────────────────────────────────────────────

def _configure_page(doc: Document) -> None:
    """Set A4 page size and GMP-standard margins."""
    for section in doc.sections:
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.0)

    # Default font for Normal style
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.styles["Normal"].font.color.rgb = BLACK


def _add_page_header(doc: Document, doc_type: str, form_data: dict) -> None:
    """Add a thin company header bar at the top of every page."""
    section = doc.sections[0]
    header  = section.header

    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run = p.add_run(
        f"PharmaGPT — Generated Document  |  {doc_type}  |  "
        f"{date.today().strftime('%d %b %Y')}"
    )
    run.font.name  = "Calibri"
    run.font.size  = Pt(8)
    run.font.color.rgb = GRAY

    # Thin border under header
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D0D9E4")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_footer(doc: Document, form_data: dict) -> None:
    """Add page number footer."""
    section = doc.sections[0]
    footer  = section.footer

    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run(f"Project: {form_data.get('project_name', '')}    |    Page ")
    run.font.name  = "Calibri"
    run.font.size  = Pt(8)
    run.font.color.rgb = GRAY

    # Auto page number field
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run2 = p.add_run()
    run2._r.append(fld_char1)
    run2._r.append(instr)
    run2._r.append(fld_char2)
    run2.font.name = "Calibri"
    run2.font.size = Pt(8)
    run2.font.color.rgb = GRAY


# ── Markdown parser ───────────────────────────────────────────────────────────

def _parse_markdown(doc: Document, content: str) -> None:
    """
    Parse markdown content line-by-line and add elements to the Document.
    Uses a simple state machine for table accumulation.
    """
    lines = content.split("\n")
    i = 0
    table_rows: list[list[str]] = []

    def flush_table():
        nonlocal table_rows
        if table_rows:
            _add_table(doc, table_rows)
            table_rows = []

    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        # ── Table detection (accumulate rows, flush on non-table line) ─────
        if stripped.startswith("|"):
            # Skip pure separator rows  |---|---|
            if re.match(r"^\|[\s\-:]+(\|[\s\-:]+)*\|?$", stripped):
                i += 1
                continue
            cells = [c.strip() for c in stripped.split("|")]
            # Remove empty first/last cells from leading/trailing |
            if cells and cells[0] == "":
                cells = cells[1:]
            if cells and cells[-1] == "":
                cells = cells[:-1]
            table_rows.append(cells)
            i += 1
            continue
        else:
            flush_table()

        # ── Headings ──────────────────────────────────────────────────────
        if stripped.startswith("#### "):
            _add_heading(doc, stripped[5:], level=3)
        elif stripped.startswith("### "):
            _add_heading(doc, stripped[4:], level=2)
        elif stripped.startswith("## "):
            _add_heading(doc, stripped[3:], level=1)
        elif stripped.startswith("# "):
            _add_title(doc, stripped[2:])

        # ── Lists ─────────────────────────────────────────────────────────
        elif stripped.startswith("- ") or stripped.startswith("* "):
            _add_list_item(doc, stripped[2:], numbered=False)
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s+", "", stripped)
            _add_list_item(doc, text, numbered=True)

        # ── Horizontal rule ───────────────────────────────────────────────
        elif stripped in ("---", "***", "___") or re.match(r"^-{3,}$", stripped):
            _add_hr(doc)

        # ── Empty line ────────────────────────────────────────────────────
        elif not stripped:
            pass   # blank lines just separate elements; no extra paragraph

        # ── Regular paragraph (with inline formatting) ────────────────────
        else:
            _add_paragraph(doc, stripped)

        i += 1

    flush_table()


# ── Element builders ──────────────────────────────────────────────────────────

def _add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(12)


def _add_heading(doc: Document, text: str, level: int) -> None:
    colors  = {1: NAVY, 2: BLUE, 3: BLACK}
    sizes   = {1: Pt(13), 2: Pt(11.5), 3: Pt(10.5)}
    spaces  = {1: Pt(16), 2: Pt(12), 3: Pt(8)}

    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name  = "Calibri"
    run.font.size  = sizes.get(level, Pt(11))
    run.font.color.rgb = colors.get(level, BLACK)
    p.paragraph_format.space_before = spaces.get(level, Pt(8))
    p.paragraph_format.space_after  = Pt(4)

    if level == 1:
        # Bottom border under major section headers
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1565C0")
        pBdr.append(bottom)
        p._p.get_or_add_pPr().append(pBdr)


def _add_paragraph(doc: Document, text: str) -> None:
    """Add a normal paragraph with basic inline **bold** and *italic* support."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    _apply_inline(p, text)


def _apply_inline(p, text: str) -> None:
    """Split text on **bold** and *italic* markers and add styled runs."""
    # Pattern: **bold**, *italic*, or plain text
    pattern = r"(\*\*[^*]+\*\*|\*[^*]+\*)"
    parts   = re.split(pattern, text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(10.5)
        elif part.startswith("*") and part.endswith("*"):
            run = p.add_run(part[1:-1])
            run.italic = True
            run.font.name = "Calibri"
            run.font.size = Pt(10.5)
        else:
            if part:
                run = p.add_run(part)
                run.font.name = "Calibri"
                run.font.size = Pt(10.5)


def _add_list_item(doc: Document, text: str, numbered: bool) -> None:
    style = "List Number" if numbered else "List Bullet"
    try:
        p = doc.add_paragraph(style=style)
    except KeyError:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
    _apply_inline(p, text)
    p.paragraph_format.space_after = Pt(2)


def _add_hr(doc: Document) -> None:
    """Add a horizontal rule paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D0D9E4")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    """Add a styled Word table from a list of cell-value rows."""
    if not rows:
        return

    # Normalise column count (some rows may be shorter)
    col_count = max(len(r) for r in rows)
    rows = [r + [""] * (col_count - len(r)) for r in rows]

    tbl = doc.add_table(rows=len(rows), cols=col_count)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"

    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.name = "Calibri"
            run.font.size = Pt(9.5)

            if r_idx == 0:
                # Header row: navy background, white bold text
                run.bold = True
                run.font.color.rgb = WHITE
                _set_cell_bg(cell, "0A2342")
            else:
                run.font.color.rgb = BLACK
                if r_idx % 2 == 0:
                    _set_cell_bg(cell, "F4F7FB")

    doc.add_paragraph()  # spacing after table


def _set_cell_bg(cell, hex_color: str) -> None:
    """Set the background colour of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)
