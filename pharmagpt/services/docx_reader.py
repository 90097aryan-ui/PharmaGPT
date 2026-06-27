"""
services/docx_reader.py — Extract text from Word (.docx) files using python-docx.

Extracts:
  - All paragraph text (body, headings, lists)
  - Table cell content, formatted as pipe-delimited rows so Gemini
    can understand tabular GMP data (acceptance criteria, test results, etc.)

DOCX files do not have a native "page count" concept (pages are calculated
by Word at render time). We estimate page count as total_words / 300,
which is a reasonable approximation for technical pharmaceutical documents.
"""

from docx import Document


def extract(file_path: str) -> tuple[str, int]:
    """
    Extract text from a .docx file and return (text, estimated_page_count).

    Parameters
    ----------
    file_path : str — absolute path to the .docx file

    Returns
    -------
    text               : str — extracted text
    estimated_pages    : int — word_count / 300 (minimum 1)
    """
    doc = Document(file_path)
    sections: list[str] = []

    # ── Paragraphs ──────────────────────────────────────────────────────────────
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Prefix headings so Gemini can identify document structure
            if para.style.name.startswith("Heading"):
                sections.append(f"\n## {text}")
            else:
                sections.append(text)

    # ── Tables ──────────────────────────────────────────────────────────────────
    # GMP documents contain critical tables (IQ checklists, OQ test results, etc.)
    for table in doc.tables:
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            sections.append("\n[Table]\n" + "\n".join(rows))

    full_text = "\n\n".join(sections)
    word_count = len(full_text.split())
    estimated_pages = max(1, word_count // 300)

    return full_text, estimated_pages
